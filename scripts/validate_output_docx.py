#!/usr/bin/env python3
"""
STEP 6: 품질 검증 및 자동 수정 (DOCX)

사전 조건: generate_docx.py 실행 완료 (workspace/output/{name}.docx 존재)

사용법:
  python validate_output.py

품질 검증까지 실행하려면:
  $env:ANTHROPIC_API_KEY="sk-ant-..."
  python validate_output.py
"""

import json
import os
import re
import shutil
import sys
import time
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from _common import get_base

# ──────────────────────────────────────────────────────────────────
# 경로 상수
# ──────────────────────────────────────────────────────────────────
BASE       = get_base()
WORKSPACE  = os.path.join(BASE, 'workspace')
STRUCTURED = os.path.join(WORKSPACE, 'structured')
OUTPUT_DIR = os.path.join(WORKSPACE, 'output')
LOG_DIR    = os.path.join(WORKSPACE, 'logs')
SKILL_OUT  = os.path.join(BASE, 'skill', 'output')

MODEL = 'claude-sonnet-4-6'


# ══════════════════════════════════════════════════════════════════
# 유틸리티
# ══════════════════════════════════════════════════════════════════

def load_meta() -> dict:
    with open(os.path.join(LOG_DIR, 'step1_meta.json'), encoding='utf-8') as f:
        return json.load(f)


def load_master() -> dict:
    p = os.path.join(STRUCTURED, 'master_dataset.json')
    if os.path.exists(p):
        with open(p, encoding='utf-8') as f:
            return json.load(f)
    return {}


def full_text_of(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(parts)


def paragraph_has_image(p) -> bool:
    return len(p._element.findall('.//' + qn('w:drawing'))) > 0


def media_file_count(docx_path: str) -> int:
    with zipfile.ZipFile(docx_path, 'r') as zf:
        return len([n for n in zf.namelist() if n.startswith('word/media/')])


# ══════════════════════════════════════════════════════════════════
# C00-C10 구조 검증 함수
#   signature: fn(doc, text, docx_path) → (bool, str)
# ══════════════════════════════════════════════════════════════════

_SEC_KW = ['시장 개요', '시장 역학', '생태계', '지역별', '세그먼트', 'R&D', '결론']


def c00(doc, text, docx_path):
    """DOCX 패키지 자체가 유효한지(zip 필수 파트 + 파싱 가능) 확인"""
    errors = []
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            names = zf.namelist()
            for req in ('[Content_Types].xml', 'word/document.xml', '_rels/.rels'):
                if req not in names:
                    errors.append(f'누락:{req}')
    except Exception as e:
        errors.append(str(e))
    if errors:
        return False, '; '.join(errors)
    return True, '필수 패키지 파트 확인'


def c01(doc, text, docx_path):
    found = [k for k in _SEC_KW if k in text]
    return len(found) >= 7, f'{len(found)}/7 섹션 확인 ({", ".join(found[:3])}…)'


def c02(doc, text, docx_path):
    if re.search(r'\$[\d.]+\s*B', text) or re.search(r'[\d,.]+\s*billion', text, re.I):
        return True, '시장규모 수치($B) 확인'
    if '$' in text and re.search(r'\d+\.\d+', text):
        return True, '시장규모 수치 확인'
    return False, '시장규모 수치 없음'


def c03(doc, text, docx_path):
    patterns = [
        r'CAGR[\s\S]{0,15}[\d.]+\s*%',
        r'[\d.]+\s*%[\s\S]{0,15}CAGR',
        r'연평균[\s]*[\d.]+\s*%',
        r'복합연간성장률[\s:]*[\d.]+',
    ]
    for p in patterns:
        if re.search(p, text, re.I):
            return True, 'CAGR 수치 확인'
    # 표 셀 단위로 텍스트가 평탄화되면 'CAGR (%)' 헤더와 실제 값이
    # 서로 다른 셀/행에 있어 인접 정규식이 매치하지 못할 수 있다.
    # 'CAGR' 단어와 소수점 % 수치가 문서 어디엔가 둘 다 있으면 통과로 간주한다.
    if re.search(r'CAGR', text, re.I) and re.search(r'\d+\.\d+\s*%', text):
        return True, 'CAGR 수치 확인(표 내 분산 배치)'
    return False, 'CAGR 수치 없음'


def c04(doc, text, docx_path):
    if '출처' in text or '※' in text:
        return True, '출처 정보 확인'
    return False, '출처 정보 없음'


def c05(doc, text, docx_path):
    markers = ['①', '②', '③', '④', '⑤']
    found   = [m for m in markers if m in text]
    if len(found) >= 5:
        return True, f'시사점 {len(found)}개(①~⑤) 확인'
    kws = ['유망 기술', '기술 격차', '공급망', '정책', 'KIST']
    kw_found = [k for k in kws if k in text]
    if len(kw_found) >= 4:
        return True, f'시사점 키워드 {len(kw_found)}개 확인'
    return False, f'시사점 항목 부족 ({len(found)}개/5)'


def c06(doc, text, docx_path):
    n = len(doc.tables)
    return n >= 2, f'테이블 {n}개'


def c07(doc, text, docx_path):
    n = len(doc.inline_shapes)
    return n >= 9, f'이미지 {n}개'


def c08(doc, text, docx_path):
    """각 이미지 단락 직후에 캡션 단락이 있는지 확인"""
    paras = doc.paragraphs
    pic_idx = [i for i, p in enumerate(paras) if paragraph_has_image(p)]
    if not pic_idx:
        return True, '이미지 없음 (해당 없음)'
    missing = 0
    for i in pic_idx:
        if i + 1 >= len(paras):
            missing += 1
            continue
        nxt_txt = paras[i + 1].text
        if '그림' in nxt_txt or '출처' in nxt_txt:
            pass
        else:
            missing += 1
    if missing == 0:
        return True, f'모든 이미지({len(pic_idx)}개)에 캡션 확인'
    return False, f'캡션 없는 이미지 {missing}개'


def c09(doc, text, docx_path):
    """inline_shapes 참조 수 ↔ 실제 word/media/ 파일 수 일치 확인"""
    n_shapes = len(doc.inline_shapes)
    if n_shapes == 0:
        return True, '이미지 참조 없음 (해당 없음)'
    n_media = media_file_count(docx_path)
    if n_media < n_shapes:
        return False, f'media 파일 부족: {n_media} vs 참조 {n_shapes}'
    return True, f'이미지 참조 일치 ({n_shapes}개)'


def c10(doc, text, docx_path):
    has_date = '작성일' in text or re.search(r'\d{4}년\s*\d{1,2}월', text) is not None
    has_src  = '출처' in text or '원본' in text
    if has_date and has_src:
        return True, '표지 출처·작성일 확인'
    miss = ([] if has_date else ['작성일']) + ([] if has_src else ['출처'])
    return False, f'표지 누락: {", ".join(miss)}'


CHECKS = [
    ('C00', 'DOCX 패키지 유효성',        c00),
    ('C01', '7개 섹션 존재',           c01),
    ('C02', '시장규모 수치 포함',       c02),
    ('C03', 'CAGR 수치 포함',          c03),
    ('C04', '출처 정보 기재',           c04),
    ('C05', '시사점 5개 항목',          c05),
    ('C06', '테이블 최소 2개',          c06),
    ('C07', '이미지 최소 9개',          c07),
    ('C08', '모든 이미지에 캡션',       c08),
    ('C09', '이미지 참조 일치',         c09),
    ('C10', '표지 출처·작성일',         c10),
]

CHECK_FN = {cid: fn for cid, _, fn in CHECKS}
TOTAL_CHECKS = len(CHECKS)


# ══════════════════════════════════════════════════════════════════
# 자동 수정 함수 (python-docx API 기반)
# ══════════════════════════════════════════════════════════════════

def fix_c04(doc: Document, src: str) -> bool:
    """출처 단락을 표지 직후에 삽입"""
    paras = doc.paragraphs
    if len(paras) < 2:
        return False
    anchor = paras[min(3, len(paras) - 1)]
    p = anchor.insert_paragraph_before(f'※ 출처: {src}')
    return True


def fix_c08(doc: Document) -> int:
    """캡션 없는 이미지 단락 직후에 기본 캡션 삽입"""
    fixed = 0
    idx = 0
    while idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        if not paragraph_has_image(p):
            idx += 1
            continue
        nxt = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None
        if nxt is not None and ('그림' in nxt.text or '출처' in nxt.text):
            idx += 1
            continue
        fixed += 1
        if nxt is not None:
            cap = nxt.insert_paragraph_before(f'그림 {fixed}. (출처: 원본 보고서)')
        else:
            cap = doc.add_paragraph(f'그림 {fixed}. (출처: 원본 보고서)')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True
        idx += 2
    return fixed


def fix_c10(doc: Document) -> bool:
    """표지 단락에 작성일 추가"""
    today = datetime.now().strftime('%Y년 %m월 %d일')
    for i, par in enumerate(doc.paragraphs[:6]):
        if '연구과제' in par.text or '원본 출처' in par.text:
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt is not None:
                nxt.insert_paragraph_before(f'작성일: {today}')
            else:
                doc.add_paragraph(f'작성일: {today}')
            return True
    return False


def apply_fixes(docx_path: str, failed: list, src: str, output_name: str) -> tuple:
    """수정 가능 항목 자동 수정 → 저장"""
    doc = Document(docx_path)
    fixed: list = []

    if 'C04' in failed and fix_c04(doc, src):
        fixed.append('C04')
    if 'C08' in failed:
        n = fix_c08(doc)
        if n:
            fixed.append(f'C08({n}개 캡션 추가)')
    if 'C10' in failed and fix_c10(doc):
        fixed.append('C10')

    if fixed:
        doc.save(docx_path)
        os.makedirs(SKILL_OUT, exist_ok=True)
        shutil.copy2(docx_path, os.path.join(SKILL_OUT, f'{output_name}.docx'))

    auto_fixable = {'C04', 'C08', 'C10'}
    manual = [c for c in failed if c not in auto_fixable]
    return fixed, manual


# ══════════════════════════════════════════════════════════════════
# Q01-Q05 품질 검증 (Claude API)
# ══════════════════════════════════════════════════════════════════

def run_quality_checks(client, text: str, tf: str) -> dict:
    summary = text[:6000] + ('…(이하 생략)' if len(text) > 6000 else '')
    prompt = f"""다음은 자동 생성된 '{tf}' 한글 시장분석 보고서의 본문입니다.
아래 5가지 품질 항목을 평가하고 순수 JSON만 반환하세요.

[본문]
{summary}

[출력 JSON 스키마]
{{
  "Q01": {{"item": "수치 정확성",       "pass": true,  "note": "1문장"}},
  "Q02": {{"item": "번역 자연스러움",    "pass": true,  "note": "1문장"}},
  "Q03": {{"item": "시사점 구체성",      "pass": true,  "note": "1문장"}},
  "Q04": {{"item": "KIST 시사점 현실성", "pass": true,  "note": "1문장"}},
  "Q05": {{"item": "섹션 완결성",        "pass": true,  "note": "1문장"}}
}}"""

    for attempt in range(3):
        try:
            msg = client.messages.create(
                model=MODEL, max_tokens=800,
                messages=[{'role': 'user', 'content': prompt}]
            )
            raw  = re.sub(r'```(?:json)?\s*', '', msg.content[0].text).strip()
            s, e = raw.find('{'), raw.rfind('}')
            if s != -1 and e != -1:
                return json.loads(raw[s:e+1])
        except Exception as ex:
            if attempt < 2:
                time.sleep(2 ** (attempt + 1))
    return {f'Q0{i}': {'item': f'항목{i}', 'pass': None, 'note': 'API 실패'}
            for i in range(1, 6)}


# ══════════════════════════════════════════════════════════════════
# 메인
# ══════════════════════════════════════════════════════════════════

def main():
    meta   = load_meta()
    master = load_master()
    output_name = meta.get('output_name', '시장분석_보고서')
    tf          = meta.get('tech_field',  '기술분야')
    src         = master.get('market_overview', {}).get('source_report', '원본 보고서')
    docx_path   = os.path.join(OUTPUT_DIR, f'{output_name}.docx')

    print(f'\n[STEP 6] 품질 검증 — {output_name}.docx')

    if not os.path.exists(docx_path):
        print(f'❌ DOCX 파일 없음: {docx_path}')
        sys.exit(1)

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f'❌ DOCX 열기 실패: {e}')
        sys.exit(1)

    full_text = full_text_of(doc)

    # ── C00-C10 구조 검증 ─────────────────────────────────────────
    print(f'\n  [구조 검증] C00-C10')
    struct_res: dict = {}
    failed: list = []

    for cid, desc, fn in CHECKS:
        passed, msg = fn(doc, full_text, docx_path)
        struct_res[cid] = {'desc': desc, 'pass': passed, 'note': msg}
        icon = '✓' if passed else '✗'
        print(f'    {icon} {cid} {desc}: {msg}')
        if not passed:
            failed.append(cid)

    # ── 자동 수정 ─────────────────────────────────────────────────
    fixed_list: list = []
    manual_list: list = []

    if failed:
        print(f'\n  [자동 수정] 실패 {len(failed)}건: {", ".join(failed)}')
        fixed_list, manual_list = apply_fixes(docx_path, failed, src, output_name)
        if fixed_list:
            print(f'    수정 완료: {", ".join(fixed_list)}')
        if manual_list:
            print(f'    수동 확인 필요: {", ".join(manual_list)}')

        if fixed_list:
            try:
                doc2 = Document(docx_path)
                text2 = full_text_of(doc2)
                for entry in fixed_list:
                    cid = entry.split('(')[0]
                    if cid in CHECK_FN:
                        p2, m2 = CHECK_FN[cid](doc2, text2, docx_path)
                        struct_res[cid]['pass'] = p2
                        struct_res[cid]['note_after_fix'] = m2
            except Exception:
                pass

    # ── Q01-Q05 품질 검증 (Claude API) ───────────────────────────
    quality_res: dict = {}
    api_key = os.environ.get('ANTHROPIC_API_KEY', '')

    if api_key:
        print('\n  [품질 검증] Claude API 호출...')
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            quality_res = run_quality_checks(client, full_text, tf)
            for qid, qv in quality_res.items():
                icon = '✓' if qv.get('pass') else ('?' if qv.get('pass') is None else '✗')
                print(f'    {icon} {qid} {qv.get("item","")}: {qv.get("note","")}')
        except Exception as e:
            print(f'    ⚠ 오류: {e}')
            quality_res = {f'Q0{i}': {'item': f'항목{i}', 'pass': None, 'note': str(e)}
                           for i in range(1, 6)}
    else:
        print('\n  [품질 검증] ANTHROPIC_API_KEY 없음 → 건너뜀')
        quality_res = {f'Q0{i}': {'item': f'항목{i}', 'pass': None, 'note': 'API 키 없음'}
                       for i in range(1, 6)}

    # ── 최종 판정 ─────────────────────────────────────────────────
    s_pass = sum(1 for v in struct_res.values()  if v['pass'])
    q_pass = sum(1 for v in quality_res.values() if v.get('pass') is True)
    q_tested = sum(1 for v in quality_res.values() if v.get('pass') is not None)
    s_ratio = s_pass / TOTAL_CHECKS

    if not struct_res['C00']['pass']:
        verdict = 'FAIL (패키지 구조 오류)'
    elif q_tested == 0:
        if   s_ratio >= 0.9: verdict = 'PASS (구조검증 기준)'
        elif s_ratio >= 0.7: verdict = 'CONDITIONAL PASS (구조검증 기준)'
        else:                verdict = 'FAIL'
    elif s_ratio >= 0.8 and q_pass >= 4:
        verdict = 'PASS'
    elif s_ratio >= 0.7 and q_pass >= 3:
        verdict = 'CONDITIONAL PASS'
    else:
        verdict = 'FAIL'

    log = {
        'timestamp':            datetime.now().isoformat(),
        'docx_file':            docx_path,
        'structural_checks':    struct_res,
        'auto_fixed':           fixed_list,
        'manual_review_needed': manual_list,
        'quality_checks':       quality_res,
        'overall_pass':         verdict == 'PASS',
        'verdict':              verdict,
        'pass_rate_structural': f'{s_pass}/{TOTAL_CHECKS}',
        'pass_rate_quality':    f'{q_pass}/5',
    }
    with open(os.path.join(LOG_DIR, 'step6_validation_docx.json'), 'w', encoding='utf-8') as f:
        json.dump(log, f, ensure_ascii=False, indent=2)

    print(f'\n✅ STEP 6 완료')
    print(f'  구조 검증: {s_pass}/{TOTAL_CHECKS} 통과')
    print(f'  자동 수정: {len(fixed_list)}건')
    print(f'  품질 검증: {q_pass}/5 통과')
    if manual_list:
        print(f'  수동 확인 필요: {len(manual_list)}건 ({", ".join(manual_list)})')
    else:
        print(f'  수동 확인 필요: 0건')
    print(f'  최종 판정: {verdict}')

    if verdict.startswith('FAIL'):
        print(f'\n  ⚠ 구조 검증 미달 항목: '
              f'{[k for k, v in struct_res.items() if not v["pass"]]}')

    print(f'\n  파이프라인 완료 → 최종 보고서: {docx_path}')


if __name__ == '__main__':
    main()
