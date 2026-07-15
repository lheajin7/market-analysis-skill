#!/usr/bin/env python3
"""
STEP 5: 한글 보고서 본문 생성 (DOCX 조립)

사전 조건: STEP 3(master_dataset.json) + generate_charts.py 실행 완료

사용법:
  python generate_docx.py

python-docx 라이브러리를 사용해 검증된 Office Open XML(.docx)을 생성한다.
"""

import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as docx_qn
from PIL import Image as PILImage

from _common import get_base, seg_title

# 개조식 문단 1개당 최대 5줄(약 200자)을 넘지 않도록 분할한다.
MAX_PARA_CHARS = 200


def _eun_neun(word: str) -> str:
    """마지막 글자의 받침 유무에 따라 '은'/'는' 조사를 고른다 — 데이터에서 가져온
    지역명·라벨을 문장에 끼워 넣을 때 조사가 하드코딩돼 있으면(예: 항상 '은') 받침
    없는 단어에 붙었을 때 어색해진다("북미은" 등). 한글 음절이 아니면(영문 라벨 등)
    안전하게 병기한다."""
    if not word:
        return '은(는)'
    code = ord(word[-1]) - 0xAC00
    if 0 <= code <= 11171:
        return '는' if code % 28 == 0 else '은'
    return '은(는)'


def _split_para(text: str, max_chars: int = MAX_PARA_CHARS) -> list:
    """5줄(약 200자) 초과 문단을 문장 경계에서 잘라 개조식 여러 항목으로 나눈다."""
    if len(text) <= max_chars:
        return [text]
    sentences = re.split(r'(?<=[.!?])\s+|(?<=함\.)\s*|(?<=음\.)\s*|(?<=됨\.)\s*', text)
    chunks, cur = [], ''
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        if cur and len(cur) + len(s) + 1 > max_chars:
            chunks.append(cur)
            cur = s
        else:
            cur = f'{cur} {s}'.strip() if cur else s
    if cur:
        chunks.append(cur)
    return chunks or [text]


# ──────────────────────────────────────────────────────────────────
# 경로 상수
# ──────────────────────────────────────────────────────────────────
BASE       = get_base()
WORKSPACE  = os.path.join(BASE, 'workspace')
STRUCTURED = os.path.join(WORKSPACE, 'structured')
CHARTS     = os.path.join(WORKSPACE, 'charts')
OUTPUT_DIR = os.path.join(WORKSPACE, 'output')
LOG_DIR    = os.path.join(WORKSPACE, 'logs')

NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
GRAY  = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)

# A4 + 표준 여백(위·아래 2.54cm, 좌·우 1.9cm) — HWPX(generate_hwpx.py)도 동일 여백을
# 쓰므로 두 포맷의 본문 폭이 거의 같아지고, 차트 폰트 크기 역산(generate_charts.py의
# DOCX_WIDTH_IN/HWPX_WIDTH_IN)도 이 값과 일치시켜야 한다 — 여백을 바꾸면 세 곳(여기,
# generate_hwpx.py, generate_charts.py) 모두 함께 갱신할 것.
PAGE_W_CM, PAGE_H_CM   = 21.0, 29.7
MARGIN_TOP_CM          = 2.54
MARGIN_BOTTOM_CM       = 2.54
MARGIN_LR_CM           = 1.9
CONTENT_WIDTH_IN = (PAGE_W_CM - 2 * MARGIN_LR_CM) / 2.54   # 본문 폭(인치) — 표·그림 공통 폭

# matplotlib/Pillow로 생성하는 차트는 본문 폭에 꽉 채워도 보통 3.5~4.5in 높이로 나오지만,
# 원본 보고서에서 재활용하는 이미지(*_image 필드)는 정사각형에 가까운 경우가 있어(예:
# 4단계 가치사슬 다이어그램 974x924px) 본문 폭에 그대로 맞추면 세로로 5.5in 이상 길어져
# 다른 그림보다 눈에 띄게 커 보인다(실측 확인, 사용자 피드백). 모든 그림에 공통으로
# 세로 높이 상한을 두어, 상한을 넘는 경우에만 폭을 비례 축소한다 — 정상 범위인 차트는
# 전혀 영향받지 않고, 종횡비가 유별난 이미지만 자동으로 작아진다.
MAX_IMAGE_HEIGHT_IN = 4.5


# ══════════════════════════════════════════════════════════════════
# DocxBuilder: 문단·이미지·표를 누적해 python-docx Document를 조립한다
# ══════════════════════════════════════════════════════════════════

class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        self._set_page()
        self._set_base_font()
        self.n_images = 0
        self.n_tables = 0
        self.fig_n = 0   # 그림 번호 자동 채번 — 호출 순서대로 1부터 증가
        self.tbl_n = 0   # 표 번호 자동 채번

    def _set_page(self):
        sec = self.doc.sections[0]
        sec.page_width   = Cm(PAGE_W_CM)
        sec.page_height  = Cm(PAGE_H_CM)
        sec.top_margin    = Cm(MARGIN_TOP_CM)
        sec.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        sec.left_margin   = Cm(MARGIN_LR_CM)
        sec.right_margin  = Cm(MARGIN_LR_CM)

    def _set_base_font(self):
        style = self.doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(10)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(docx_qn('w:rFonts'))
        if rfonts is None:
            rfonts = rpr.makeelement(docx_qn('w:rFonts'), {})
            rpr.append(rfonts)
        rfonts.set(docx_qn('w:eastAsia'), '맑은 고딕')

    # ── 공개 단락 메서드 ─────────────────────────────────────────
    def h1(self, text: str):
        h = self.doc.add_heading('', level=1)
        run = h.add_run(text)
        run.font.color.rgb = NAVY
        run.font.name = '맑은 고딕'

    def h2(self, text: str):
        h = self.doc.add_heading('', level=2)
        run = h.add_run(text)
        run.font.color.rgb = NAVY
        run.font.name = '맑은 고딕'

    def normal(self, text: str, bold: bool = False):
        for chunk in _split_para(text):
            p = self.doc.add_paragraph()
            run = p.add_run(chunk)
            run.bold = bold
            run.font.color.rgb = BLACK

    def bullet(self, text: str):
        for chunk in _split_para(text):
            self.doc.add_paragraph(chunk, style='List Bullet')

    def caption(self, text: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.italic = True

    def empty(self):
        self.doc.add_paragraph()

    # ── 이미지 ───────────────────────────────────────────────────
    def image(self, path: str, title: str, desc: str = ''):
        """표(table)와 동일하게 본문 폭(CONTENT_WIDTH_IN)에 꽉 채워 삽입한다.
        캡션은 "그림 N. 제목" 형태로 자동 채번한다 — 호출부에서 번호를 직접
        하드코딩하지 않으므로 그림 순서가 바뀌어도 번호가 어긋나지 않는다. 출처는
        표지(첫 페이지 요약)에만 싣고 본문 그림·표 하단에는 반복하지 않는다는
        원칙에 따라 캡션에 "(출처: ...)"를 붙이지 않는다.
        desc는 그 그림의 실제 내용(구체적 수치·경향 등)을 담고 있을 때만 넘기고,
        제목을 다른 말로 반복하는 문장이면 아예 넘기지 않는다(빈 문자열은 생략됨)."""
        if not os.path.exists(path):
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        with PILImage.open(path) as im:
            w_px, h_px = im.size
        height_at_full_width = CONTENT_WIDTH_IN * h_px / w_px if w_px else 0
        if height_at_full_width > MAX_IMAGE_HEIGHT_IN:
            run.add_picture(path, height=Inches(MAX_IMAGE_HEIGHT_IN))
        else:
            run.add_picture(path, width=Inches(CONTENT_WIDTH_IN))
        self.n_images += 1
        self.fig_n += 1
        self.caption(f'그림 {self.fig_n}. {title}')
        if desc:
            self.normal(desc)

    # ── 표 ───────────────────────────────────────────────────────
    def table(self, headers: list, rows: list, title: str = '', desc: str = '', unit: str = ''):
        """이미지와 동일하게 본문 폭(CONTENT_WIDTH_IN)에 꽉 채워 고정 폭으로 삽입한다.
        (python-docx는 열 폭을 지정해도 tblLayout이 auto면 Word가 창 크기에 맞춰
        다시 늘려버리므로, tblLayout을 fixed로 강제하고 tblW/각 열·셀 폭을 명시한다.)
        title을 넘기면 "표 N. 제목" 캡션을 그림과 동일한 규칙으로 자동 채번해 표
        아래에 붙인다(출처는 표지에만 싣고 여기서는 반복하지 않는다).
        unit을 넘기면 "(단위: {unit})"을 표 바로 위 우측 정렬로 한 번만 표시한다 —
        금액·시장규모처럼 모든 셀에 같은 단위가 반복되는 표는 각 셀에 단위를 넣지
        말고 이 인자로 대표 표시할 것(원칙 — 아래 참고)."""
        if not headers or not rows:
            return
        if unit:
            up = self.doc.add_paragraph()
            up.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            urun = up.add_run(f'(단위: {unit})')
            urun.font.size = Pt(9)
            urun.font.color.rgb = GRAY
        n_col = len(headers)
        all_rows = [headers] + rows
        t = self.doc.add_table(rows=len(all_rows), cols=n_col)
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False   # python-docx가 tblLayout을 fixed로 맞춰준다

        # add_table()이 기본으로 넣어주는 <w:tblW type="auto" w="0"/>을 그대로 두면
        # Word가 창 크기에 맞춰 다시 늘여버리므로, 새로 추가하지 말고 기존 요소를
        # 찾아 값을 덮어써야 한다 (그냥 append하면 tblW가 중복 삽입되어 문서가
        # 손상된 것으로 인식될 수 있다).
        tblPr = t._tbl.tblPr
        tblW = tblPr.find(docx_qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(docx_qn('w:type'), 'dxa')
        tblW.set(docx_qn('w:w'), str(round(CONTENT_WIDTH_IN * 1440)))

        col_width = Inches(CONTENT_WIDTH_IN / n_col)
        for col in t.columns:
            col.width = col_width

        for ri, row_data in enumerate(all_rows):
            for ci in range(n_col):
                text = str(row_data[ci]) if ci < len(row_data) and row_data[ci] is not None else ''
                cell = t.cell(ri, ci)
                cell.width = col_width
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
        self.n_tables += 1
        if title:
            self.tbl_n += 1
            self.caption(f'표 {self.tbl_n}. {title}')
        if desc:
            self.normal(desc)
        self.doc.add_paragraph()  # 표 뒤 여백

    def save(self, path: str):
        self.doc.save(path)


# ══════════════════════════════════════════════════════════════════
# 데이터 로드
# ══════════════════════════════════════════════════════════════════

def load_data():
    with open(os.path.join(STRUCTURED, 'master_dataset.json'), encoding='utf-8') as f:
        md = json.load(f)
    with open(os.path.join(CHARTS, 'chart_catalog.json'), encoding='utf-8') as f:
        cc = json.load(f)
    return md, cc


def get_chart_map(catalog: dict) -> dict:
    return {c['chart_id']: c['filename']
            for c in catalog.get('charts', [])
            if c['status'] in ('generated', 'reused')}


def resolve_path(chart_id: str, chart_map: dict) -> str | None:
    fname = chart_map.get(chart_id)
    if not fname:
        return None
    path = os.path.join(CHARTS, fname)
    return path if os.path.exists(path) else None


# ══════════════════════════════════════════════════════════════════
# 보고서 본문 조립
# ══════════════════════════════════════════════════════════════════

def build_report(b: DocxBuilder, md: dict, chart_map: dict, meta: dict):
    tf   = md.get('tech_field',    '기술분야')
    proj = md.get('project_name',  '연구과제')
    sec1 = md.get('market_overview',   {})
    sec2 = md.get('market_dynamics',   {})
    sec3 = md.get('ecosystem',         {})
    sec4 = md.get('regional_analysis', {})
    sec5 = md.get('segmentation',      {})
    sec6 = md.get('rnd_trends',        {})
    sec7 = md.get('implications',      {})
    src  = sec1.get('source_report', '원본 보고서')
    today = datetime.now().strftime('%Y년 %m월 %d일')

    def img(cid: str, title: str, desc: str = ''):
        path = resolve_path(cid, chart_map)
        if path:
            b.image(path, title, desc=desc)

    # ── 표지 ──────────────────────────────────────────────────────
    # 출처는 여기(표지, 첫 페이지 요약)에만 싣는다 — 본문·그림·표 하단에는
    # 반복하지 않는다는 원칙에 따라 이 아래로는 "(출처: ...)"를 붙이지 않는다.
    b.empty()
    b.h1(f'{tf} 글로벌 시장 분석 요약')
    b.normal(f'연구과제: {proj}')
    b.normal(f'원본 출처: {src}  ({meta.get("publish_year", "")})')
    b.normal(f'작성일: {today}')
    b.empty()

    # ── 섹션 1: 시장 개요 및 성장 전망 ───────────────────────────
    b.h1('1. 시장 개요 및 성장 전망')

    bg = sec1.get('growth_background', '')
    if bg:
        b.normal(bg)
    b.empty()

    pd_ = sec1.get('power_demand_data', {})
    pd_desc = ''
    if pd_.get('years') and pd_.get('demand_incl_ai_twh'):
        y0, y1 = pd_['years'][0], pd_['years'][-1]
        v0, v1 = pd_['demand_incl_ai_twh'][0], pd_['demand_incl_ai_twh'][-1]
        pd_desc = f'AI 포함 시나리오 기준 전력 수요는 {y0}년 {v0}TWh에서 {y1}년 {v1}TWh로 증가할 전망임.'
    img('V1_1', f'{tf} 분야 전력 수요 전망', desc=pd_desc)

    sc      = sec1.get('scenarios', {})
    base_s  = sec1.get('base_year_size_usd_b', '')
    base_yr = sec1.get('base_year', '')
    fc_yr   = sec1.get('forecast_year', '')
    cagr    = sec1.get('cagr_pct', '')
    opt_cagr  = sc.get('optimistic',  {}).get('cagr_pct')
    pess_cagr = sc.get('pessimistic', {}).get('cagr_pct')
    scenario_desc = ''
    if isinstance(opt_cagr, (int, float)) and isinstance(pess_cagr, (int, float)):
        diff = round(opt_cagr - pess_cagr, 1)
        scenario_desc = (f'낙관 시나리오는 연평균 {opt_cagr}%로 비관 시나리오({pess_cagr}%) '
                          f'대비 {diff}%p 높은 성장률을 전망함.')
    # 원문에 시나리오 분석이 없으면(scenarios 비어 있음) 표를 아예 만들지 않는다 —
    # 예전에는 base_year 값만 채우고 예측연도·CAGR을 '?'로 남겨 오해를 유발했다(실측 확인).
    if sc:
        tbl_hdr = ['시나리오', f'기준연도({base_yr})', f'예측연도({fc_yr})', 'CAGR (%)']
        tbl_row = [
            ['낙관 (Optimistic)',
             f'{base_s}' if base_s else '-',
             f'{sc.get("optimistic",  {}).get("size_usd_b", "?")}',
             f'{sc.get("optimistic",  {}).get("cagr_pct", "?")}%'],
            ['현실 (Realistic)',
             f'{base_s}' if base_s else '-',
             f'{sc.get("realistic",  {}).get("size_usd_b", "?")}',
             f'{sc.get("realistic",  {}).get("cagr_pct", cagr)}%'],
            ['비관 (Pessimistic)',
             f'{base_s}' if base_s else '-',
             f'{sc.get("pessimistic", {}).get("size_usd_b", "?")}',
             f'{sc.get("pessimistic", {}).get("cagr_pct", "?")}%'],
        ]
        b.table(tbl_hdr, tbl_row, title='시나리오별 시장규모 전망', desc=scenario_desc,
                unit=sec1.get('currency_unit', '$B'))
    img('V1_2', '시나리오별 시장규모 전망')

    trends = sec1.get('key_trends', [])
    if trends:
        b.h2('핵심 트렌드')
        for t in trends:
            b.bullet(f'{t.get("trend","")}: {t.get("description","")}')

    # ── 섹션 2: 시장 역학 ─────────────────────────────────────────
    b.h1('2. 시장 역학 (Market Dynamics)')
    b.normal(f'{tf} 시장의 성장 동인, 저해 요인, 기회 요소를 분석함.')
    img('V2_1', '시장 역학 구조')

    def dyn_section(title: str, items: list):
        if items:
            b.h2(title)
            for it in items:
                b.bullet(f'{it.get("title","")}: {it.get("description","")}')

    dyn_section('성장 동인 (Drivers)',       sec2.get('drivers',      []))
    dyn_section('저해 요인 (Restraints)',    sec2.get('restraints',   []))
    if sec2.get('has_challenges'):
        dyn_section('도전 과제 (Challenges)', sec2.get('challenges', []))
    dyn_section('비즈니스 기회 (Opportunities)', sec2.get('opportunities', []))

    # ── 섹션 3: 기술·솔루션 생태계 ───────────────────────────────
    # (1) 공급망 → (2) 가치사슬 → (3) 생태계 맵 순으로 다룬다. 각 항목은 원문에
    # 실제 그림이 있으면 그 그림을 그대로 쓰고(generate_charts.py의 reuse 로직),
    # 없을 때만 합성 다이어그램으로 대체한다. desc는 STEP 3가 실제 내용을 채운
    # 경우에만 나오고, 없으면 img()가 조용히 캡션만 남긴다.
    b.h1('3. 기술·솔루션 생태계')

    # 원문에 해당 그림(재활용 이미지) 또는 합성용 구조화 데이터가 전혀 없으면
    # img()가 조용히 아무것도 그리지 않아 소제목만 남고 그 아래가 완전히 비는
    # 문제가 있었다(실측 확인) — V2_1의 "빈 카테고리는 만들지 않는다" 원칙과
    # 동일하게, 실제로 그릴 내용이 있을 때만 소제목을 낸다.
    if sec3.get('supply_chain_image') or sec3.get('supply_chain'):
        b.h2('공급망 (Supply Chain) 분석')
        img('V3_3', '공급망 구조', desc=sec3.get('supply_chain_desc', ''))

    if sec3.get('value_chain_image') or sec3.get('value_chain'):
        b.h2('가치사슬 (Value Chain) 분석')
        img('V3_2', '가치사슬 구조', desc=sec3.get('value_chain_desc', ''))

    if sec3.get('ecosystem_map_image') or sec3.get('key_players_by_category'):
        b.h2('생태계 맵 (Ecosystem Map)')
        img('V3_1', '생태계 맵', desc=sec3.get('ecosystem_map_desc', ''))

    ms = sec3.get('market_share_table', [])
    if ms:
        b.h2('경쟁구도 (Competitive Landscape)')
        top3 = ms[:3]
        top3_txt = ', '.join(f"{r.get('company','')}({r.get('share_range','')})" for r in top3)
        n_frag = max(0, len(ms) - 3)
        # 점유율 지표명·해석 문장은 데이터에서 덮어쓸 수 있다 — 기본 해석문은 "다수의 소규모
        # 기업이 1% 미만을 나눠 갖는 파편화 시장"을 전제하므로, 소수 기업만 존재하는 과점
        # 시장 데이터에 그대로 쓰면 사실과 다른 문장이 된다(generate_hwpx.py와 동일).
        default_desc = (f'상위 3개 기업({top3_txt})이 점유율 상위권을 형성하며, 나머지 {n_frag}여 개 '
                        f'기업 대부분은 1% 미만의 점유율을 보여 파편화된 경쟁구조를 나타냄. 소수 '
                        f'선도기업의 시장 지배력과 다수 틈새 기술기업 간 경쟁이 공존하는 구조임.')
        b.table(['기업명', sec3.get('market_share_metric', '시장점유율 범위')],
                [[r.get('company', ''), r.get('share_range', '')] for r in ms[:10]],
                title='경쟁구도 (Competitive Landscape)',
                desc=sec3.get('market_share_desc') or default_desc)

    nm = sec3.get('notable_movements', '')
    if nm and nm != '원문 미확인':
        b.h2('주요 시장 동향')
        b.normal(nm)

    # ── 섹션 4: 지역별 시장 분석 ─────────────────────────────────
    b.h1('4. 지역별 시장 분석')
    lead = sec4.get('leading_region', '')
    fast = sec4.get('fastest_growing_region', '')
    if lead and lead != '원문 미확인':
        b.normal(f'{lead}이(가) 시장을 주도하며, {fast}에서 가장 높은 성장률이 전망됨.')

    regs = sec4.get('regions', [])
    lead_region = next((r for r in regs if r.get('name') == lead), None)
    v4_1_desc = ''
    yrs4 = sec4.get('years', [])
    if lead_region and lead_region.get('sizes_usd_b') and len(yrs4) >= 2 \
            and len(lead_region['sizes_usd_b']) >= 2:
        v4_1_desc = (f'{lead}{_eun_neun(lead)} {yrs4[0]}년 ${lead_region["sizes_usd_b"][0]}B에서 '
                      f'{yrs4[-1]}년 ${lead_region["sizes_usd_b"][-1]}B 규모로 성장할 전망임.')
    img('V4_1', '지역별 시장규모', desc=v4_1_desc)

    if regs:
        b.h2('권역별 기회 및 리스크')
        r_hdr = ['권역', '단기 기회', '단기 리스크', '장기 기회', '장기 리스크']
        r_rows = [
            [r.get('name', '')[:20],
             r.get('short_term_opportunity', '')[:55],
             r.get('short_term_risk',        '')[:55],
             r.get('long_term_opportunity',  '')[:55],
             r.get('long_term_risk',         '')[:55]]
            for r in regs[:5]
        ]
        b.table(r_hdr, r_rows, title='권역별 기회 및 리스크')

    # V4_2(매트릭스 그림)는 삽입하지 않는다 — 바로 위 표와 완전히 같은 데이터를
    # 그림으로만 다시 보여주는 중복이었고, 표 형태의 데이터는 이미지가 아니라
    # 표로 넣는 것이 원칙이다(SKILL.md STEP 4 참고).

    # ── 섹션 5: 세그먼트별 시장 분석 ─────────────────────────────
    b.h1('5. 세그먼트별 시장 분석')

    def seg_section(h2: str, cid: str, seg_data: dict):
        b.h2(h2)
        notes = '; '.join(
            s.get('growth_note', '') for s in seg_data.get('segments', [])
            if s.get('growth_note')
        )
        if notes:
            b.normal(notes)
        img(cid, f'{h2}')

    # 세그먼트 분류 축은 보고서마다 다르므로 축 이름·개수를 데이터에서 그대로 가져온다
    # (하드코딩하면 다른 보고서를 넣었을 때 엉뚱한 축 이름이 표시된다).
    MAIN_SLOTS = ['V5_1_end_use', 'V5_2_dc_type', 'V5_3_solution', 'V5_4_rack_density']
    for i, cid in enumerate(MAIN_SLOTS):
        axis = sec5.get('axes', [])[i] if i < len(sec5.get('axes', [])) else {}
        if axis.get('segments'):
            seg_section(seg_title(axis, f'세그먼트{i+1}'), cid, axis)

    DETAIL_SLOTS = ['V5_5_air_sub', 'V5_6_liquid_dir', 'V5_7_immersion']
    for i, cid in enumerate(DETAIL_SLOTS):
        detail_axes = sec5.get('detail_axes', [])
        axis = detail_axes[i] if i < len(detail_axes) else {}
        if axis.get('segments'):
            seg_section(seg_title(axis, f'세부분류{i+1}'), cid, axis)

    # ── 섹션 6: R&D 및 기술 동향 ─────────────────────────────────
    b.h1('6. R&D 및 기술 동향')
    at = sec5.get('adoption_trend', {})
    adopt_series = at.get('series', [])
    adopt_desc = ''
    if adopt_series and adopt_series[0].get('values_pct') and at.get('years'):
        v0, v1 = adopt_series[0]['values_pct'][0], adopt_series[0]['values_pct'][-1]
        y0, y1 = at['years'][0], at['years'][-1]
        alabel = adopt_series[0].get("label", "채택률")
        adopt_desc = (f'{alabel}{_eun_neun(alabel)} {y0}년 {v0}%에서 '
                       f'{y1}년 {v1}%로 상승할 전망임.')
    img('V6_1', f'{at.get("label", "기술 채택률")} 전망', desc=adopt_desc)

    pt = sec6.get('patent_trend', {})
    if pt.get('description') and pt['description'] != '원문 미확인':
        # 절 제목도 지표에 맞춰 바뀔 수 있다(generate_hwpx.py와 동일) — 원문이 특허가 아닌
        # 다른 지표를 제시하면 section_title로 덮어쓴다(기본값은 특허 기준).
        b.h2(pt.get('section_title', '특허 동향'))
        b.normal(pt['description'])
        # 지표 라벨/단위는 데이터에 따라 가변 — 기본값은 특허 기준, 다른 지표면 스키마 필드로 덮어쓴다.
        metric_label = pt.get('metric_label', '특허 출원 건수')
        metric_unit  = pt.get('metric_unit', '건')
        # 표는 '국가별', 그림(V6_2)은 '기업별'로 분리해 분석한다 — 특허 동향은 국가
        # 단위 경쟁 구도(표)와 기업 단위 경쟁 구도(그림)를 함께 보여줄 때 가장
        # 유용하기 때문. 국가별 수치가 없으면 표는 생략하고(없는 수치를 지어내지
        # 않음) 국가명 목록은 아래 불릿으로 계속 표시된다.
        country_detail = pt.get('top_countries_detail', [])
        if country_detail:
            b.table(['국가', metric_label],
                    [[d.get('country', ''), f"{d.get('count', 0):,}{metric_unit}"] for d in country_detail[:6]],
                    title=f'국가별 {metric_label}')
        img('V6_2', f'기업별 {metric_label}',
            desc='가로축은 기업 간 큰 격차를 비교하기 쉽도록 로그 스케일로 표시함.')
        if pt.get('top_countries'):
            b.bullet(pt.get('top_countries_label', '주요 특허 보유국') + ': ' + ', '.join(pt['top_countries']))
        if pt.get('top_companies'):
            b.bullet(pt.get('top_companies_label', '주요 특허 보유 기업') + ': ' + ', '.join(pt['top_companies']))

    cases = sec6.get('case_studies', [])
    if cases:
        b.h2('케이스 스터디')
        for c in cases[:8]:
            b.bullet(f'[{c.get("organization","")}] {c.get("title","")} '
                     f'— {c.get("technology","")}: {c.get("outcome","")}')

    gov = sec6.get('government_initiatives', [])
    if gov:
        b.h2('정부 정책·이니셔티브')
        for g in gov[:5]:
            b.bullet(f'[{g.get("region","")}] {g.get("initiative","")}: '
                     f'{g.get("description","")}')

    emerging = sec6.get('emerging_technologies', [])
    if emerging:
        b.h2('신기술 동향')
        for e in emerging[:5]:
            b.bullet(f'{e.get("tech","")}: {e.get("description","")}')

    # ── 섹션 7: 결론 및 기술개발 시사점 ──────────────────────────
    b.h1('7. 결론 및 기술개발 시사점')

    b.h2('7-1. 종합 결론')
    for conc in sec7.get('key_conclusions', []):
        b.bullet(conc)

    b.h2('7-2. 기술개발 시사점')
    impl = sec7.get('implications', {})
    impl_items = [
        ('① 유망 기술 방향',         impl.get('i1_promising_tech', '')),
        ('② 기술 격차 및 추격 전략',  impl.get('i2_tech_gap',       '')),
        ('③ 공급망 참여 전략',        impl.get('i3_supply_chain',   '')),
        ('④ 정책·제도 설계 방향',     impl.get('i4_policy',         '')),
        ('⑤ KIST 연구전략 관점',      impl.get('i5_kist',           '')),
    ]
    for num_title, text in impl_items:
        if text and text != '원문 미확인':
            b.normal(num_title, bold=True)
            b.normal(text)
            b.empty()

    note = sec7.get('data_basis_note', '')
    if note:
        b.normal(f'※ {note}')


# ══════════════════════════════════════════════════════════════════
# 메인
# ══════════════════════════════════════════════════════════════════

def main():
    master_p  = os.path.join(STRUCTURED, 'master_dataset.json')
    catalog_p = os.path.join(CHARTS, 'chart_catalog.json')
    meta_p    = os.path.join(LOG_DIR,  'step1_meta.json')

    for p, name in [(master_p, 'master_dataset.json'),
                    (catalog_p, 'chart_catalog.json'),
                    (meta_p,    'step1_meta.json')]:
        if not os.path.exists(p):
            print(f'{name} 없음 — 이전 단계를 먼저 실행하세요.')
            sys.exit(1)

    md, catalog = load_data()
    with open(meta_p, encoding='utf-8') as f:
        meta = json.load(f)

    output_name = meta.get('output_name', '시장분석_보고서')
    tf = md.get('tech_field', '기술분야')
    print(f'\n[STEP 5] DOCX 보고서 생성 — {tf}')
    print(f'  출력 파일: {output_name}.docx')

    chart_map = get_chart_map(catalog)

    print('  문서 조립 중...')
    builder = DocxBuilder()
    build_report(builder, md, chart_map, meta)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, f'{output_name}.docx')
    builder.save(docx_path)
    size_kb = os.path.getsize(docx_path) // 1024

    # 섹션 포함 여부 확인 (본문 텍스트 스캔)
    full_text = '\n'.join(p.text for p in builder.doc.paragraphs)
    sec_keys  = ['시장 개요', '시장 역학', '생태계', '지역별', '세그먼트', 'R&D', '결론']
    secs_ok   = [k for k in sec_keys if k in full_text]
    secs_miss = [k for k in sec_keys if k not in full_text]

    # 최종 출력 파일을 skill/output 에도 복사
    skill_out = os.path.join(BASE, 'skill', 'output')
    os.makedirs(skill_out, exist_ok=True)
    import shutil
    shutil.copy2(docx_path, os.path.join(skill_out, f'{output_name}.docx'))

    # 로그
    log = {
        'timestamp':     datetime.now().isoformat(),
        'output_file':   docx_path,
        'size_kb':       size_kb,
        'images':        builder.n_images,
        'tables':        builder.n_tables,
        'sections_ok':   secs_ok,
        'sections_miss': secs_miss,
        'step5_status':  'success' if len(secs_ok) >= 7 else 'partial',
    }
    with open(os.path.join(LOG_DIR, 'step5_docx.json'), 'w', encoding='utf-8') as f:
        json.dump(log, f, ensure_ascii=False, indent=2)

    # 완료 출력
    all_sec = 'all' if not secs_miss else f'{", ".join(secs_miss)} 누락'
    print(f'\n✅ STEP 5 완료')
    print(f'  생성 파일: {output_name}.docx')
    print(f'  파일 크기: {size_kb} KB')
    print(f'  삽입 이미지: {builder.n_images}개')
    print(f'  삽입 표: {builder.n_tables}개')
    print(f'  7개 섹션: {"모두 포함" if not secs_miss else all_sec}')
    print(f'\n  → 다음 단계: python validate_output.py')


if __name__ == '__main__':
    main()
